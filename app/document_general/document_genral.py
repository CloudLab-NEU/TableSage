"""
用于生成最后的报告文档
pydocx
输入的数据有什么？
输出一个doc文档

需要自定义这个报告文档的格式，内容有什么？

1.展示问题的基本信息（问题描述、表格数据等）
2.展示当前问题处理的sql_skeleton,question_skeleton
3.展示当前匹配的相似问题以及其对应的相似度，以问题（相似度）形式展示
4.展示当前的第一次答案的情况（问题+模型答案+真实答案）
5.如果置信度判断结束需要第二次指导答题，则展示指导答题的选择的策略
6.最后就是目标问题的答案展示
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

class TableSageReportGenerator:
    """
    TableSage答题报告生成器
    用于生成包含完整答题过程的Word文档报告
    """
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        styles = self.doc.styles

        if 'CustomTitle' not in [style.name for style in styles]:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = '微软雅黑'
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)

        if 'CustomHeading2' not in [style.name for style in styles]:
            heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_font = heading2_style.font
            heading2_font.name = '微软雅黑'
            heading2_font.size = Pt(14)
            heading2_font.bold = True
            heading2_style.paragraph_format.space_before = Pt(12)
            heading2_style.paragraph_format.space_after = Pt(6)
        
        if 'CustomNormal' not in [style.name for style in styles]:
            normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
            normal_font = normal_style.font
            normal_font.name = '宋体'
            normal_font.size = Pt(11)
            normal_style.paragraph_format.line_spacing = 1.15
    
    def generate_report(self, process_data, output_path=None):
        """
        生成完整的答题报告
        
        Args:
            process_data (dict): 包含答题过程所有数据的字典
            output_path (str): 输出文件路径，如果为None则自动生成
        
        Returns:
            str: 生成的文档路径
        """
        self._add_document_title()

        self._add_question_info(process_data)

        self._add_skeleton_info(process_data)

        self._add_similar_questions(process_data)

        self._add_first_answer_result(process_data)

        self._add_guidance_strategy(process_data)

        self._add_final_answer(process_data)

        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"TableSage_Report_{timestamp}.docx"
        
        self.doc.save(output_path)
        return output_path
    
    def _add_document_title(self):
        """添加文档标题"""
        title = self.doc.add_paragraph("TableSage 智能问答系统报告", style='CustomTitle')

        time_para = self.doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph("=" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_question_info(self, process_data):
        """添加问题基本信息"""
        self.doc.add_paragraph("1. 问题基本信息", style='CustomHeading2')

        question = process_data.get('user_question', '未提供问题')
        self.doc.add_paragraph(f"问题描述: {question}", style='CustomNormal')

        table_data = process_data.get('user_table', {})
        if table_data:
            self.doc.add_paragraph("表格数据:", style='CustomNormal')
            
            headers = table_data.get('header', [])
            rows = table_data.get('rows', [])
            
            if headers and rows:
                display_rows = rows[:5] 
                table = self.doc.add_table(rows=len(display_rows) + 1, cols=len(headers))
                table.style = 'Table Grid'
                
                for i, header in enumerate(headers):
                    table.cell(0, i).text = str(header)
                    table.cell(0, i).paragraphs[0].runs[0].font.bold = True
                
                for row_idx, row in enumerate(display_rows):
                    for col_idx, cell_value in enumerate(row):
                        if col_idx < len(headers):
                            table.cell(row_idx + 1, col_idx).text = str(cell_value)
                
                if len(rows) > 5:
                    self.doc.add_paragraph(f"注: 表格共{len(rows)}行数据，此处仅显示前5行", style='CustomNormal')
    
    def _add_skeleton_info(self, process_data):
        """添加SQL和问题骨架信息"""
        self.doc.add_paragraph("2. 问题处理骨架", style='CustomHeading2')

        # SQL骨架
        sql_skeleton = process_data.get('sql_skeleton', '未生成SQL骨架')
        self.doc.add_paragraph("SQL骨架:", style='CustomNormal')
        sql_para = self.doc.add_paragraph(sql_skeleton)
        sql_para.style.font.name = 'Consolas'
        sql_para.style.font.size = Pt(10)
        
        # 问题骨架
        question_skeleton = process_data.get('question_skeleton', '未生成问题骨架')
        self.doc.add_paragraph("问题骨架:", style='CustomNormal')
        self.doc.add_paragraph(question_skeleton, style='CustomNormal')
    
    def _add_similar_questions(self, process_data):
        """添加相似问题匹配结果"""
        self.doc.add_paragraph("3. 相似问题匹配结果", style='CustomHeading2')
        
        similar_questions = process_data.get('similar_questions', [])
        if similar_questions:
            # 检查similar_questions的格式
            for i, sq in enumerate(similar_questions[:10], 1):  # 显示前10个
                if isinstance(sq, str):
                    try:
                        from db.db_manager import DatabaseManager
                        db_manager = DatabaseManager()
                        knowledge_record = db_manager.get_knowledge_by_id(sq)
                        if knowledge_record:
                            question_content = knowledge_record.get("question", "")
                            if question_content:
                                self.doc.add_paragraph(f"{i}. {question_content} [ID: {sq}]", style='CustomNormal')
                            else:
                                self.doc.add_paragraph(f"{i}. Table ID: {sq}", style='CustomNormal')
                        else:
                            self.doc.add_paragraph(f"{i}. Table ID: {sq}", style='CustomNormal')
                    except Exception as e:
                        self.doc.add_paragraph(f"{i}. Table ID: {sq}", style='CustomNormal')
                        
                elif isinstance(sq, dict):
                    # 如果是字典格式
                    similarity = sq.get('similarity', 0)
                    question = sq.get('question', '未知问题')
                    table_id = sq.get('table_id', sq.get('id', '未知ID'))
                    if similarity > 0:
                        self.doc.add_paragraph(f"{i}. {question} (相似度: {similarity:.3f}) [ID: {table_id}]", style='CustomNormal')
                    else:
                        self.doc.add_paragraph(f"{i}. {question} [ID: {table_id}]", style='CustomNormal')
                else:
                    # 其他格式
                    self.doc.add_paragraph(f"{i}. {str(sq)}", style='CustomNormal')
        else:
            self.doc.add_paragraph("未找到相似问题", style='CustomNormal')
    
    def _add_first_answer_result(self, process_data):
        self.doc.add_paragraph("4. 第一次答题结果", style='CustomHeading2')
        
        answer_result = process_data.get('first_answer_results', {})
        
        # 置信度
        confidence = answer_result.get('confidence', 0)
        self.doc.add_paragraph(f"置信度: {confidence:.3f}", style='CustomNormal')
        
        # 答题统计
        total_count = answer_result.get('total_count', 0)
        correct_count = answer_result.get('flag_0_count', 0)
        self.doc.add_paragraph(f"答题统计: {correct_count}/{total_count} (正确率: {correct_count/total_count*100 if total_count > 0 else 0:.1f}%)", style='CustomNormal')
        
        # 是否需要策略指导
        need_strategy = answer_result.get('need_strategy', False)
        self.doc.add_paragraph(f"是否需要策略指导: {'是' if need_strategy else '否'}", style='CustomNormal')
    
    def _add_guidance_strategy(self, process_data):
        guidance_result = process_data.get('guidance_result')
        if guidance_result:
            self.doc.add_paragraph("5. 指导答题策略", style='CustomHeading2')

            initial_confidence = guidance_result.get('initial_confidence')
            if initial_confidence is not None:
                self.doc.add_paragraph(f"初始置信度: {initial_confidence:.3f}", style='CustomNormal')
            
            recalc_confidence = guidance_result.get('recalculated_confidence')
            if recalc_confidence is not None:
                self.doc.add_paragraph(f"重新计算置信度: {recalc_confidence:.3f}", style='CustomNormal')
            
            guided_results = guidance_result.get('guided_results', [])
            if guided_results:
                self.doc.add_paragraph("指导答题详细结果:", style='CustomNormal')
                
                for i, guided_item in enumerate(guided_results, 1):

                    table_id = guided_item.get('table_id', f'问题{i}')
                    self.doc.add_paragraph(f"  {i}. 表格ID: {table_id}", style='CustomNormal')
                    
                    strategies_tried = guided_item.get('strategies_tried', [])
                    if strategies_tried:
                        strategies_text = " → ".join(strategies_tried)
                        self.doc.add_paragraph(f"     策略尝试: {strategies_text}", style='CustomNormal')

                    is_correct = guided_item.get('is_correct', False)
                    result_text = "✓ 正确" if is_correct else "✗ 错误"
                    self.doc.add_paragraph(f"     指导结果: {result_text}", style='CustomNormal')

                    model_answer = guided_item.get('model_answer', '')
                    if model_answer:
                        self.doc.add_paragraph(f"     模型答案: {model_answer}", style='CustomNormal')

                    true_answer = guided_item.get('true_answer', '')
                    if true_answer:
                        if isinstance(true_answer, list):
                            true_answer_text = str(true_answer)
                        else:
                            true_answer_text = str(true_answer)
                        self.doc.add_paragraph(f"     真实答案: {true_answer_text}", style='CustomNormal')
                    
                    if i < len(guided_results):
                        self.doc.add_paragraph("     " + "-" * 30, style='CustomNormal')

            if guided_results:
                total_guided = len(guided_results)
                correct_guided = sum(1 for item in guided_results if item.get('is_correct', False))
                self.doc.add_paragraph(f"指导答题统计: {correct_guided}/{total_guided} (成功率: {correct_guided/total_guided*100:.1f}%)", style='CustomNormal')
    
    def _add_final_answer(self, process_data):
        if process_data.get('guidance_result'):
            self.doc.add_paragraph("6. 最终答案", style='CustomHeading2')
        else:
            self.doc.add_paragraph("5. 最终答案", style='CustomHeading2')
        
        final_answer = process_data.get('final_answer', '未获取最终答案')
        self.doc.add_paragraph(f"最终答案: {final_answer}", style='CustomNormal')
        
        flow_path = process_data.get('flow_path', 'unknown')
        flow_desc = {
            'direct': '直接答题（置信度足够）',
            'guidance': '指导答题（置信度不足）',
            'error': '错误处理'
        }
        self.doc.add_paragraph(f"答题流程: {flow_desc.get(flow_path, flow_path)}", style='CustomNormal')

def create_process_data_from_result(user_question, user_table, tablesage_result):
    """
    从TableSageProcessor的结果创建用于文档生成的数据结构
    
    Args:
        user_question (str): 用户问题
        user_table (dict): 用户表格
        tablesage_result (dict): TableSageProcessor.process()的返回结果
        answer_result (dict): 第一次答题结果
        guidance_result (dict): 指导答题结果
        true_answer (str): 真实答案
    
    Returns:
        dict: 格式化的处理数据
    """
    process_data = {
        'user_question': user_question, 
        'user_table': user_table,
        'sql_skeleton': tablesage_result.get('sql_skeleton', ''),
        'question_skeleton': tablesage_result.get('question_skeleton', ''),
        'similar_questions': tablesage_result.get('similar_questions', []),
        'final_answer': tablesage_result.get('answer', ''),
        'confidence': tablesage_result.get('confidence', 0),
        'first_answer_results': tablesage_result.get('first_answer_results', {}),
        'gudiance_result': tablesage_result.get('guidance_result', {}),
        'flow_path': tablesage_result.get('flow_path', 'unknown'),
    }

    if tablesage_result.get('guidance_result'):
        process_data['guidance_result'] = tablesage_result['guidance_result']
    
    return process_data

def generate_tablesage_report(user_question, user_table, processor_result, 
                            output_dir="./reports", true_answer=None):
    """
    生成TableSage答题报告的便捷函数
    
    Args:
        user_question (str): 用户问题
        user_table (dict): 用户表格
        processor_result (dict): TableSageProcessor的完整结果
        output_dir (str): 输出目录
    
    Returns:
        str: 生成的报告文件路径
    """
    os.makedirs(output_dir, exist_ok=True)
    
    generator = TableSageReportGenerator()

    process_data = create_process_data_from_result(
        user_question, user_table, processor_result
    )

    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = os.path.join(output_dir, f"TableSage_Report_{timestamp}.docx")
    
    return generator.generate_report(process_data, output_path)